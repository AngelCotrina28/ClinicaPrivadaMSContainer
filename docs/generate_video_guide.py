from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from generate_report import (
    ARCHITECTURE_IMAGE,
    BLUE,
    DATABASE_IMAGE,
    GREEN,
    TEAL,
    VALIDATION_IMAGE,
    add_bullet,
    add_docx_table,
    add_numbered,
    configure_docx,
)


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "docs" / "entregable"
OUTPUT.mkdir(parents=True, exist_ok=True)
GUIDE_PATH = OUTPUT / "Guia-Tecnica-y-Guion-Video-Microservicios.docx"


def add_paragraph(document: Document, text: str, bold_start: str | None = None) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if bold_start and text.startswith(bold_start):
        paragraph.add_run(bold_start).bold = True
        paragraph.add_run(text[len(bold_start):])
    else:
        paragraph.add_run(text)


def add_code(document: Document, code: str) -> None:
    table = document.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    cell.text = ""
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), "F4F6F8")
    cell._tc.get_or_add_tcPr().append(shading)
    for index, line in enumerate(code.strip("\n").splitlines()):
        paragraph = cell.paragraphs[0] if index == 0 else cell.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(0)
        run = paragraph.add_run(line)
        run.font.name = "Consolas"
        run.font.size = Pt(8.5)
        run.font.color.rgb = RGBColor.from_string("263238")
    document.add_paragraph()


def add_callout(document: Document, title: str, text: str, color: str = "D9EAF7") -> None:
    table = document.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), color)
    cell._tc.get_or_add_tcPr().append(shading)
    paragraph = cell.paragraphs[0]
    paragraph.add_run(f"{title}: ").bold = True
    paragraph.add_run(text)
    document.add_paragraph()


def add_toc(document: Document) -> None:
    paragraph = document.add_paragraph()
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = 'TOC \\o "1-3" \\h \\z \\u'
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    placeholder = OxmlElement("w:t")
    placeholder.text = "Actualice la tabla de contenido con clic derecho > Actualizar campo."
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction, separate, placeholder, end])


def add_endpoint_table(document: Document, rows: list[list[str]]) -> None:
    add_docx_table(document, ["Metodo", "Ruta por Gateway", "Funcion"], rows, [2.1, 7.6, 7.1])


def add_section_break(document: Document) -> None:
    document.add_page_break()


def create_guide() -> None:
    document = Document()
    configure_docx(document)

    for _ in range(2):
        document.add_paragraph()
    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("GUIA TECNICA, ONBOARDING Y GUION DE VIDEO\nARQUITECTURA DE MICROSERVICIOS")
    run.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor.from_string(BLUE)
    subtitle = document.add_paragraph("Proyecto Luz del Tunel - Clinica Privada")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].bold = True
    subtitle.runs[0].font.size = Pt(15)
    document.add_paragraph()
    add_docx_table(document, ["Dato", "Contenido"], [
        ["Curso", "Desarrollo de Sistemas Web"],
        ["Grupo", "11"],
        ["Equipo", "Cotrina Ruiz, Angel Eduardo; Moya Laureano, Jussel Enrique; Tocto Caqui, Juan Andres"],
        ["Cuentas GitHub", "AngelCotrina28; Enrike352; RexoJuan"],
        ["Objetivo", "Explicar implementacion, codigo, bases de datos, despliegue, pruebas y funcionamiento"],
        ["Estado documentado", "13 de julio de 2026"],
    ], [4.2, 12.5])
    document.add_picture(str(ARCHITECTURE_IMAGE), width=Inches(6.7))
    document.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_callout(
        document,
        "Como usar esta guia",
        "La primera mitad explica el sistema como material de incorporacion tecnica. La segunda mitad funciona como libreto del video y banco de preguntas para la sustentacion.",
    )
    add_section_break(document)

    document.add_heading("Indice", level=1)
    add_toc(document)
    add_section_break(document)

    document.add_heading("1. Resumen ejecutivo", level=1)
    add_paragraph(
        document,
        "El proyecto partio de una aplicacion Angular conectada a un backend Spring Boot centralizado y una base MySQL. Para cumplir el entregable de microservicios se extrajeron capacidades con limites de negocio claros, se incorporaron servicios de infraestructura y se mantuvo una ruta de compatibilidad hacia el backend original. El resultado es una migracion gradual: el frontend conoce una sola entrada, el API Gateway, mientras cada servicio nuevo posee su propio codigo, puerto y base de datos."
    )
    add_paragraph(
        document,
        "La solucion contiene tres microservicios de negocio obligatorios: Citas, Atencion Medica y Caja/Facturacion. Utilizan MongoDB, PostgreSQL y MySQL respectivamente. Se agregaron Auth y Notificaciones, ademas de Gateway, Config Server y Eureka. El conjunto se puede ejecutar mediante Docker Compose o Kubernetes local, y los servicios principales tambien cuentan con configuracion para Render y bases administradas en Aiven o MongoDB Atlas."
    )
    add_callout(
        document,
        "Idea principal para la exposicion",
        "No dividimos el sistema por carpetas arbitrarias. Lo dividimos por responsabilidades del negocio, propiedad de datos y capacidad de despliegue independiente.",
        "E2F0D9",
    )

    document.add_heading("1.1. Que se solicito", level=2)
    for text in [
        "Un minimo de tres microservicios de negocio.",
        "Bases de datos independientes usando tres motores distintos.",
        "Servicios de infraestructura adicionales.",
        "Despliegue local con Docker Compose y Kubernetes.",
        "Ajuste del frontend para utilizar la infraestructura de microservicios.",
        "Video con camara y audio mostrando implementacion, despliegue y funcionamiento.",
    ]:
        add_bullet(document, text)

    document.add_heading("1.2. Que se implemento", level=2)
    add_docx_table(document, ["Tipo", "Componente", "Resultado"], [
        ["Negocio", "Citas Service", "Reservas, disponibilidad, validacion de cruces y cancelacion; MongoDB."],
        ["Negocio", "Atencion Medica Service", "Registro, consulta, cierre y anulacion de atenciones; PostgreSQL."],
        ["Negocio", "Caja Facturacion Service", "Deudas, pagos, saldos, comprobantes y resumen; MySQL."],
        ["Apoyo", "Notificaciones Service", "Avisos, estado de lectura y resumen; PostgreSQL."],
        ["Infraestructura", "Auth Service", "Usuarios, roles, login, BCrypt y JWT; MySQL."],
        ["Infraestructura", "API Gateway", "Entrada unica, proxy HTTP, rutas nuevas y fallback temporal."],
        ["Infraestructura", "Config Server", "Configuracion centralizada para el entorno local."],
        ["Infraestructura", "Eureka Server", "Registro y descubrimiento de servicios en local."],
    ], [3.4, 5.2, 8])

    add_section_break(document)
    document.add_heading("2. Veredicto del punto 4: frontend a microservicios", level=1)
    add_callout(
        document,
        "Respuesta corta",
        "Si esta implementado en el codigo local, pero todavia no esta cerrado como entrega porque los cambios no tienen commit/push y la aplicacion publica de Vercel aun debe redeplegarse con esta version.",
        "FCE4D6",
    )
    add_docx_table(document, ["Comprobacion", "Estado", "Evidencia"], [
        ["Angular usa una URL unica", "Implementado", "environment.ts conserva apiUrl: /api."],
        ["Desarrollo local usa Gateway", "Implementado", "proxy.conf.json reenvia /api a http://localhost:8090."],
        ["Vercel usa Gateway cloud", "Implementado localmente", "vercel.json reescribe /api hacia clinica-gateway-service.onrender.com."],
        ["Gateway conoce los servicios", "Implementado", "AUTH_SERVICE_URL, CITAS_SERVICE_URL, ATENCION_SERVICE_URL, CAJA_SERVICE_URL y NOTIFICACIONES_SERVICE_URL."],
        ["Cambios publicados en GitHub", "Pendiente", "Frontend y Container presentan cambios sin commit."],
        ["Frontend publico redeplegado", "Pendiente", "Vercel debe construir despues del push."],
        ["Gateway cloud redeplegado", "Pendiente", "Render debe construir la nueva ruta /api/ms y los flags de compatibilidad."],
    ], [5.2, 3.4, 8])

    document.add_heading("2.1. Cambio exacto en Angular", level=2)
    add_paragraph(document, "Antes, el proxy local enviaba directamente al backend original en el puerto 8080. Ahora envia al Gateway en 8090:")
    add_code(document, '''{
  "/api": {
    "target": "http://localhost:8090",
    "secure": false,
    "changeOrigin": true
  }
}''')
    add_paragraph(document, "En produccion, Vercel mantiene /api para Angular y lo reescribe al Gateway de Render:")
    add_code(document, '''{
  "source": "/api/:path*",
  "destination": "https://clinica-gateway-service.onrender.com/api/:path*"
}''')
    add_paragraph(
        document,
        "Esta decision evita colocar cinco URLs distintas dentro de componentes Angular. Todos los servicios frontend siguen construyendo rutas con environment.apiUrl; la infraestructura decide el destino real."
    )

    document.add_heading("2.2. Migracion gradual y prefijo /api/ms", level=2)
    add_paragraph(
        document,
        "Los contratos del backend original y los microservicios nuevos no son identicos en todos los modulos. Enviar todas las rutas antiguas directamente a los servicios nuevos romperia pantallas existentes. Por eso el Gateway admite rutas explicitas bajo /api/ms y conserva un fallback al backend original. Este enfoque es una variante del patron Strangler: se reemplazan capacidades por etapas."
    )
    add_docx_table(document, ["Ruta recibida", "Destino"], [
        ["/api/auth/**", "Auth Service"],
        ["/api/notificaciones/**", "Notificaciones Service"],
        ["/api/ms/citas/**", "Citas Service, removiendo /ms antes de reenviar"],
        ["/api/ms/atenciones/**", "Atencion Medica Service"],
        ["/api/ms/caja/**", "Caja Facturacion Service"],
        ["/api/** no migrada", "Backend Spring Boot original"],
    ], [7.5, 9.2])
    add_callout(
        document,
        "Respuesta para un senior",
        "El punto 4 esta resuelto a nivel de entrada arquitectonica. La migracion funcional completa de cada pantalla es una etapa posterior y esta representada explicitamente por el fallback, no ocultada como si ya estuviera terminada.",
        "E2F0D9",
    )

    add_section_break(document)
    document.add_heading("3. Evolucion de la arquitectura", level=1)
    document.add_heading("3.1. Situacion inicial", level=2)
    add_paragraph(
        document,
        "La aplicacion inicial tenia una arquitectura cliente-servidor clasica: Angular consumia /api, Vercel o el proxy local enviaban las solicitudes al backend Spring Boot, y este concentraba autenticacion, citas, atenciones, farmacia, caja y administracion sobre una base MySQL centralizada. Esa version sigue siendo util porque contiene modulos que aun no fueron extraidos."
    )
    add_code(document, "Frontend Angular -> Backend Spring Boot -> MySQL central")

    document.add_heading("3.2. Arquitectura objetivo implementada", level=2)
    document.add_picture(str(ARCHITECTURE_IMAGE), width=Inches(6.7))
    document.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_paragraph(
        document,
        "El frontend ya no decide a que backend llamar. El Gateway es la frontera de entrada. Auth, Citas, Atencion, Caja y Notificaciones poseen almacenamiento separado. Config Server y Eureka apoyan el entorno local; Render usa URLs configuradas por variables y puede mantenerlos deshabilitados."
    )

    document.add_heading("3.3. Por que estos limites", level=2)
    add_docx_table(document, ["Limite", "Razon funcional", "Razon tecnica"], [
        ["Citas", "Agenda y disponibilidad cambian con frecuencia.", "Documento flexible, consulta por medico/fecha e indice de estado."],
        ["Atencion Medica", "Contiene diagnostico, tratamiento e historial.", "Datos estructurados y consistentes en PostgreSQL."],
        ["Caja Facturacion", "Maneja dinero, saldo y pago.", "Transacciones y precision decimal en MySQL."],
        ["Notificaciones", "Es transversal a varios flujos.", "No debe contaminar la logica de citas o caja."],
        ["Auth", "Identidad y permisos son transversales.", "JWT permite autenticacion stateless entre servicios."],
    ], [4.1, 6.2, 6.4])

    add_section_break(document)
    document.add_heading("4. Distribucion del codigo", level=1)
    add_paragraph(
        document,
        "La solucion se conserva en tres repositorios principales. Esta separacion permite desplegar frontend, backend heredado y contenedor de microservicios de manera independiente."
    )
    add_docx_table(document, ["Repositorio", "Responsabilidad", "Elementos principales"], [
        ["Clinica-Privada-Frontend", "Interfaz Angular y sesion del usuario.", "src/app, environment.ts, proxy.conf.json, vercel.json."],
        ["Clinica-Privada-Backend", "Backend original y fallback temporal.", "Modulos clinicos no migrados, API REST y MySQL central."],
        ["Clinica-Privada-Microservices-Container", "Servicios independientes y orquestacion.", "8 proyectos Maven, Docker Compose, k8s, database, postman y scripts."],
    ], [5.2, 5.5, 6])

    document.add_heading("4.1. Arbol del contenedor", level=2)
    add_code(document, '''Clinica-Privada-Microservices-Container/
  Clinica-Config-Server/
  Clinica-Eureka-Server/
  Clinica-Auth-Service/
  Clinica-Citas-Service/
  Clinica-Atencion-Medica-Service/
  Clinica-Caja-Facturacion-Service/
  Clinica-Notificaciones-Service/
  Clinica-Gateway-Service/
  database/
  postman/
  scripts/
  k8s/
  docker-compose.microservices.yml
  pom.xml''')
    add_paragraph(
        document,
        "El pom.xml raiz es un agregador Maven. No mezcla el codigo de los servicios; permite ejecutar una compilacion de reactor sobre los ocho modulos. Cada servicio conserva su propio pom.xml, Dockerfile y application.properties."
    )

    document.add_heading("4.2. Capas internas de un microservicio", level=2)
    add_docx_table(document, ["Capa", "Responsabilidad", "Ejemplo"], [
        ["controllers", "Expone endpoints HTTP y codigos de respuesta.", "CitaController, AtencionMedicaController."],
        ["dtos", "Define contratos de entrada y salida; aplica validaciones.", "CrearCitaRequestDTO, PagoResponseDTO."],
        ["services", "Contiene reglas de negocio y transacciones.", "CitaService, CajaFacturacionService."],
        ["entities", "Representa la persistencia y estados del dominio.", "Cita, AtencionMedica, Deuda, Pago."],
        ["repositories", "Abstrae consultas JPA o MongoDB.", "CitaRepository, DeudaRepository."],
        ["config", "Configura seguridad, JWT, CORS y filtros.", "SecurityConfig, JwtAuthenticationFilter."],
    ], [3.2, 7.1, 6.4])

    document.add_heading("4.3. Codigo del Gateway", level=2)
    add_docx_table(document, ["Archivo", "Que hace"], [
        ["GatewayProxyController.java", "Captura todas las solicitudes /api/**."],
        ["GatewayRouter.java", "Resuelve el destino segun prefijo y flags de migracion."],
        ["ProxyService.java", "Copia cuerpo, metodo, query y headers permitidos; ejecuta la llamada HTTP."],
        ["GatewayProperties.java", "Carga URLs y flags desde application.properties o variables de entorno."],
        ["GatewayInfoController.java", "Publica /gateway/routes para inspeccionar el mapa activo."],
    ], [7.2, 9.5])
    add_code(document, '''String microservicePath = removerPrefijoMicroservicio(path);
if (microservicePath != null) {
    return resolverMicroservicioExplicito(microservicePath);
}
return properties.getBackendUrl() + path;''')

    add_section_break(document)
    document.add_heading("5. Microservicios en detalle", level=1)
    document.add_heading("5.1. Auth Service", level=2)
    add_paragraph(
        document,
        "Auth Service es responsable de identidad. Guarda roles y usuarios en clinica_auth_db, cifra contrasenas con BCrypt y emite JWT firmados con HS256. El token contiene el username como subject, el rol como claim y una fecha de expiracion."
    )
    add_endpoint_table(document, [
        ["POST", "/api/auth/login", "Valida credenciales y devuelve token, usuario y rol."],
        ["GET", "/api/auth/me", "Devuelve el perfil del usuario autenticado."],
        ["GET", "/api/auth/roles", "Lista roles."],
        ["GET", "/api/auth/usuarios", "Lista usuarios; requiere ADMINISTRADOR."],
        ["POST", "/api/auth/usuarios", "Crea usuario; requiere ADMINISTRADOR."],
    ])
    add_paragraph(
        document,
        "AuthSeedService crea los roles ADMINISTRADOR, RECEPCIONISTA, MEDICO, FARMACIA y CAJERO. Tambien crea un administrador inicial cuando la base esta vacia, usando variables de entorno y no una contrasena fija dentro del codigo."
    )

    document.add_heading("5.2. Citas Service", level=2)
    add_paragraph(
        document,
        "Citas usa MongoDB y almacena cada cita como documento. La regla mas importante es evitar que un medico tenga horarios superpuestos. CitaService valida que la hora final sea posterior, que el bloque pertenezca al horario de atencion y que no exista otra cita activa que se cruce."
    )
    add_endpoint_table(document, [
        ["POST", "/api/ms/citas", "Crea una cita y bloquea cruces de horario."],
        ["GET", "/api/ms/citas/{id}", "Obtiene una cita."],
        ["GET", "/api/ms/citas/paciente/{id}", "Lista citas del paciente."],
        ["GET", "/api/ms/citas/medico/{id}?fecha=", "Lista agenda del medico por fecha."],
        ["GET", "/api/ms/citas/disponibilidad", "Calcula disponibilidad del medico."],
        ["PATCH", "/api/ms/citas/{id}/cancelar", "Cancela y registra el motivo."],
    ])

    document.add_heading("5.3. Atencion Medica Service", level=2)
    add_paragraph(
        document,
        "Atencion Medica usa PostgreSQL. Relaciona historia clinica, paciente, medico y opcionalmente una cita. Registra motivo, diagnostico, tratamiento e indicaciones. Su ciclo de estado es REGISTRADA, CERRADA o ANULADA."
    )
    add_endpoint_table(document, [
        ["POST", "/api/ms/atenciones", "Registra una atencion."],
        ["GET", "/api/ms/atenciones/{id}", "Consulta una atencion."],
        ["GET", "/api/ms/atenciones/historia/{id}", "Lista por historia clinica."],
        ["GET", "/api/ms/atenciones/paciente/{id}", "Lista por paciente."],
        ["GET", "/api/ms/atenciones/medico/{id}", "Lista por medico."],
        ["PATCH", "/api/ms/atenciones/{id}/cerrar", "Completa tratamiento e indicaciones y cierra."],
        ["PATCH", "/api/ms/atenciones/{id}/anular", "Anula una atencion."],
    ])

    document.add_heading("5.4. Caja Facturacion Service", level=2)
    add_paragraph(
        document,
        "Caja usa MySQL y separa Deuda de Pago. El saldo se deriva de monto_total menos monto_pagado. El servicio impide pagar una deuda anulada o exceder el saldo. Cada pago genera un numero de comprobante y actualiza el estado a PAGADA cuando el saldo llega a cero."
    )
    add_endpoint_table(document, [
        ["POST", "/api/ms/caja/deudas", "Crea una deuda."],
        ["GET", "/api/ms/caja/deudas", "Filtra deudas por paciente o estado."],
        ["GET", "/api/ms/caja/deudas/{id}", "Consulta deuda y saldo."],
        ["POST", "/api/ms/caja/pagos", "Registra un pago."],
        ["GET", "/api/ms/caja/pagos/{id}", "Consulta un pago."],
        ["GET", "/api/ms/caja/pagos/paciente/{id}", "Lista pagos del paciente."],
        ["GET", "/api/ms/caja/resumen", "Calcula indicadores agregados de caja."],
    ])

    document.add_heading("5.5. Notificaciones Service", level=2)
    add_paragraph(
        document,
        "Notificaciones es un servicio de apoyo porque no representa el nucleo de una atencion, pero puede ser consumido por varios dominios. Guarda destinatario, tipo, canal, mensaje, referencia y fechas del ciclo de envio/lectura."
    )
    add_endpoint_table(document, [
        ["POST", "/api/ms/notificaciones", "Crea una notificacion."],
        ["GET", "/api/ms/notificaciones", "Lista notificaciones."],
        ["GET", "/api/ms/notificaciones/destinatario/{tipo}/{id}", "Lista avisos de un destinatario."],
        ["GET", "/api/ms/notificaciones/resumen", "Devuelve conteos por estado."],
        ["PATCH", "/api/ms/notificaciones/{id}/leida", "Marca una notificacion como leida."],
    ])

    document.add_heading("5.6. Config Server y Eureka", level=2)
    add_paragraph(
        document,
        "Config Server expone propiedades comunes desde Clinica-Config-Server/config-repo. Eureka permite que las aplicaciones se registren y puedan ser observadas en localhost:8761. En Docker Compose y Kubernetes se levantan antes que los servicios de negocio. En Render se deshabilitan porque cada URL se configura directamente mediante variables de entorno; esto reduce consumo en el plan academico."
    )

    add_section_break(document)
    document.add_heading("6. Bases de datos y propiedad de datos", level=1)
    document.add_picture(str(DATABASE_IMAGE), width=Inches(6.7))
    document.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_paragraph(
        document,
        "Compartir motor no significa compartir base. Auth y Caja utilizan MySQL, pero auth escribe clinica_auth_db y caja escribe clinica_caja_db. Atencion y Notificaciones usan PostgreSQL en bases distintas. Citas posee clinica_citas_db en MongoDB. Ningun servicio debe consultar tablas o colecciones de otro servicio directamente."
    )
    add_docx_table(document, ["Servicio", "Motor", "Base", "Datos propios"], [
        ["Auth", "MySQL", "clinica_auth_db", "roles, usuarios"],
        ["Citas", "MongoDB", "clinica_citas_db", "citas e indices"],
        ["Atencion", "PostgreSQL", "clinica_atencion_db", "atenciones_medicas"],
        ["Caja", "MySQL", "clinica_caja_db", "deudas, pagos"],
        ["Notificaciones", "PostgreSQL", "clinica_notificaciones_db", "notificaciones"],
    ], [3.6, 3.2, 5, 4.8])
    add_paragraph(
        document,
        "La carpeta database/ versiona schema.sql, data.sql y scripts MongoDB. Hibernate puede crear o actualizar tablas en el arranque, pero los scripts explicitos permiten revisar, documentar y reproducir el modelo sin depender solo del ORM."
    )
    add_callout(
        document,
        "Tradeoff",
        "No existe una transaccion ACID unica que abarque MongoDB, PostgreSQL y MySQL. El flujo de demostracion coordina llamadas secuenciales. En una evolucion real se usarian eventos, idempotencia y una Saga para compensar fallos parciales.",
        "FCE4D6",
    )

    add_section_break(document)
    document.add_heading("7. Flujo completo de una solicitud", level=1)
    document.add_heading("7.1. Login y propagacion del JWT", level=2)
    for text in [
        "Angular envia POST /api/auth/login al mismo origen.",
        "El proxy de Angular o Vercel reenvia /api al Gateway.",
        "GatewayRouter selecciona Auth Service.",
        "Auth valida username y BCrypt contra clinica_auth_db.",
        "Auth devuelve un JWT con username, rol, emision y expiracion.",
        "Angular guarda el token y jwtInterceptor agrega Authorization: Bearer <token>.",
        "Gateway copia el header Authorization al servicio destino.",
        "Cada microservicio valida firma y expiracion con el mismo JWT_SECRET.",
    ]:
        add_numbered(document, text)
    add_code(document, "Authorization: Bearer eyJhbGciOiJIUzI1NiJ9...")

    document.add_heading("7.2. Flujo clinico demostrable", level=2)
    add_docx_table(document, ["Paso", "Solicitud", "Resultado"], [
        ["1", "POST /api/ms/citas", "Cita PROGRAMADA en MongoDB."],
        ["2", "POST /api/ms/atenciones", "Atencion REGISTRADA en PostgreSQL."],
        ["3", "PATCH /api/ms/atenciones/{id}/cerrar", "Atencion CERRADA."],
        ["4", "POST /api/ms/caja/deudas", "Deuda PENDIENTE en MySQL."],
        ["5", "POST /api/ms/caja/pagos", "Pago y comprobante; deuda PAGADA."],
        ["6", "POST /api/ms/notificaciones", "Aviso de pago en PostgreSQL."],
        ["7", "PATCH /api/ms/citas/{id}/cancelar", "Limpieza de la cita de prueba."],
    ], [2, 7.5, 7.2])
    add_paragraph(
        document,
        "scripts/verify-microservices.ps1 y la coleccion Postman automatizan este flujo, guardando los IDs que devuelve cada respuesta para utilizarlos en el paso siguiente."
    )

    add_section_break(document)
    document.add_heading("8. Seguridad implementada", level=1)
    add_docx_table(document, ["Control", "Implementacion", "Alcance"], [
        ["Contrasenas", "BCryptPasswordEncoder", "No se comparan ni almacenan como texto plano."],
        ["Sesion", "JWT stateless", "No depende de sesion HTTP en memoria."],
        ["Firma", "HS256 con secreto Base64", "Todos los servicios validan el mismo emisor logico."],
        ["Autorizacion", "Spring Security", "Usuarios requiere rol ADMINISTRADOR; otros endpoints exigen autenticacion."],
        ["Transporte interno", "Gateway copia Authorization", "El token llega al servicio final."],
        ["CORS", "Lista de origenes por variable", "Permite Angular local o Vercel segun ambiente."],
        ["Kubernetes", "Secret y ConfigMap", "Separa credenciales de configuracion no sensible."],
    ], [3.5, 5.8, 7.4])
    add_callout(
        document,
        "Limitacion conocida",
        "Los servicios comparten un secreto simetrico JWT. Para produccion se recomendaria firma asimetrica, rotacion de claves, un proveedor de identidad y politicas de autorizacion mas finas.",
        "FCE4D6",
    )

    add_section_break(document)
    document.add_heading("9. Despliegue local con Docker Compose", level=1)
    add_paragraph(
        document,
        "Docker construye una imagen por aplicacion y crea contenedores para bases, infraestructura y servicios. Compose expresa el orden y las dependencias de salud. Las bases usan volumenes para conservar datos entre reinicios."
    )
    add_code(document, '''cd "C:\\Users\\Angel\\Desktop\\DWEB CASO CLINICA\\Clinica-Privada-Microservices-Container"
.\\scripts\\start-microservices.ps1
.\\scripts\\verify-microservices.ps1''')
    add_docx_table(document, ["Fase", "Que ocurre"], [
        ["1. Bases", "MySQL Auth, MySQL Caja, PostgreSQL Atencion, PostgreSQL Notificaciones y MongoDB Citas."],
        ["2. Config", "Config Server queda saludable en 8888."],
        ["3. Registro", "Eureka inicia en 8761."],
        ["4. Servicios", "Auth 8091, Notificaciones 8092, Citas 8093, Atencion 8094 y Caja 8095."],
        ["5. Entrada", "Gateway publica 8090."],
        ["6. Prueba", "El script ejecuta login y flujo clinico completo."],
    ], [4.2, 12.5])
    add_paragraph(document, "Para detener el entorno:")
    add_code(document, ".\\scripts\\stop-microservices.ps1")
    add_callout(
        document,
        "Regla operativa",
        "No mantener Docker Compose y Kubernetes levantados al mismo tiempo. Duplican bases y servicios, consumen memoria y vuelven confuso que instancia responde.",
        "FCE4D6",
    )

    add_section_break(document)
    document.add_heading("10. Despliegue local con Kubernetes", level=1)
    add_paragraph(
        document,
        "Kubernetes no reemplaza las imagenes Docker: las ejecuta y administra. Los manifiestos declaran el estado deseado. Docker Desktop aporta el cluster local docker-desktop y el runtime de contenedores."
    )
    add_docx_table(document, ["Manifiesto", "Contenido"], [
        ["00-namespace-config.yaml", "Namespace clinica-ms, Secret y ConfigMap compartidos."],
        ["01-databases.yaml", "PVC, Deployments y Services de cinco bases."],
        ["02-config-eureka.yaml", "Config Server y Eureka."],
        ["03-microservices.yaml", "Deployments y Services de Auth, Notificaciones y los tres servicios de negocio."],
        ["04-gateway.yaml", "Gateway, URLs internas y Service."],
    ], [6.4, 10.3])
    add_paragraph(document, "Cada Deployment tiene una readinessProbe. Kubernetes no marca el Pod como disponible hasta que la base responda o /actuator/health devuelva salud. Las bases utilizan estrategia Recreate para evitar dos Pods escribiendo el mismo volumen local.")
    add_code(document, '''.\\scripts\\start-k8s.ps1
.\\scripts\\status-k8s.ps1
.\\scripts\\verify-k8s.ps1''')
    add_paragraph(document, "Para exponer manualmente el Gateway:")
    add_code(document, "kubectl -n clinica-ms port-forward svc/gateway-service 8090:8090")
    add_paragraph(document, "Para detener y eliminar el namespace:")
    add_code(document, ".\\scripts\\stop-k8s.ps1")

    document.add_heading("10.1. Que mostrar en Kubernetes", level=2)
    for text in [
        "kubectl config current-context debe indicar docker-desktop.",
        "kubectl -n clinica-ms get pods debe mostrar todos los Pods Running y READY 1/1.",
        "kubectl -n clinica-ms get svc debe mostrar resolucion interna por nombre.",
        "kubectl -n clinica-ms get pvc debe mostrar volumenes Bound.",
        "Eureka debe listar las aplicaciones registradas.",
        "verify-k8s.ps1 debe terminar con Validacion completa: OK.",
    ]:
        add_bullet(document, text)

    add_section_break(document)
    document.add_heading("11. Despliegue en nube", level=1)
    add_paragraph(
        document,
        "Render aloja los backends. Aiven aloja MySQL y PostgreSQL. MongoDB Atlas aloja la base de Citas. Vercel aloja Angular. En nube, Config Server y Eureka estan deshabilitados para los servicios de Render; las rutas se entregan por variables de entorno. Esta es una decision de costo y simplicidad para la demostracion, no una afirmacion de que Eureka sea obligatorio en cualquier nube."
    )
    add_docx_table(document, ["Componente", "URL o proveedor"], [
        ["Frontend", "https://clinica-privada-frontend.vercel.app/"],
        ["Gateway", "https://clinica-gateway-service.onrender.com"],
        ["Auth", "https://clinica-auth-service.onrender.com"],
        ["Citas", "https://clinica-citas-service.onrender.com"],
        ["Atencion", "https://clinica-atencion-medica-service.onrender.com"],
        ["Caja", "https://clinica-caja-facturacion-service.onrender.com"],
        ["Notificaciones", "https://clinica-notificaciones-service.onrender.com"],
        ["Bases", "Aiven MySQL/PostgreSQL y MongoDB Atlas"],
    ], [5.2, 11.5])
    add_callout(
        document,
        "Antes de grabar",
        "Los servicios gratuitos de Render pueden estar dormidos. Abrir /actuator/health, esperar status UP y ejecutar Postman antes de iniciar la grabacion.",
        "FCE4D6",
    )

    add_section_break(document)
    document.add_heading("12. Pruebas y evidencias", level=1)
    document.add_picture(str(VALIDATION_IMAGE), width=Inches(6.7))
    document.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_docx_table(document, ["Prueba", "Comando o artefacto", "Que demuestra"], [
        ["Reactor Maven", ".\\mvnw.cmd test", "Los ocho modulos compilan; Gateway ejecuta pruebas unitarias."],
        ["Gateway Router", "GatewayRouterTest", "Alias /api/ms, fallback y Auth se resuelven al destino esperado."],
        ["Angular", "npm.cmd run build", "El frontend compila con apiUrl=/api y salida Vercel valida."],
        ["Docker/K8s", "verify-microservices.ps1", "Flujo real a traves del Gateway y bases."],
        ["Postman", "postman/*.json", "Coleccion importable para local y Render."],
        ["Esquemas", "database/", "DDL, datos de ejemplo e indices MongoDB versionados."],
    ], [4, 5.8, 6.9])
    add_callout(
        document,
        "Lectura honesta de las pruebas",
        "Actualmente las pruebas unitarias nuevas cubren el enrutamiento del Gateway. Los demas servicios compilan y se validan mediante pruebas integrales de API, pero no poseen una suite unitaria amplia. No conviene afirmar cobertura total.",
        "FCE4D6",
    )

    add_section_break(document)
    document.add_heading("13. Distribucion del trabajo y dominio para exponer", level=1)
    add_paragraph(
        document,
        "El documento anterior define los roles de proyecto: Angel como Desarrollador Principal, Jussel como Scrum Master y Juan como Product Owner; los tres participan como Full Stack. El repositorio de microservicios concentra actualmente sus commits en AngelCotrina28. Para una exposicion tecnica creible, cada integrante debe dominar y demostrar una parte, sin atribuir commits inexistentes."
    )
    add_docx_table(document, ["Integrante", "Rol", "Bloque recomendado para dominar"], [
        ["Angel Eduardo Cotrina", "Desarrollador Principal", "Evolucion arquitectonica, Gateway, Auth, ajuste Angular y decisiones de compatibilidad."],
        ["Jussel Enrique Moya", "Scrum Master", "Tres servicios de negocio, motores, propiedad de datos y flujo Cita-Atencion-Caja."],
        ["Juan Andres Tocto", "Product Owner", "Docker Compose, Kubernetes, nube, demostracion Postman y valor para la clinica."],
    ], [4.5, 4, 8.2])
    add_paragraph(
        document,
        "Esta distribucion es para estudio, mantenimiento y presentacion. Antes de grabar, cada integrante debe ejecutar personalmente los comandos de su bloque y ser capaz de ubicar los archivos mencionados."
    )

    add_section_break(document)
    document.add_heading("14. Guion del video con camara y audio", level=1)
    add_callout(
        document,
        "Duracion sugerida",
        "Entre 20 y 28 minutos. Mantener la camara visible, compartir pantalla con letra legible y evitar mostrar contrasenas o archivos .env.",
        "E2F0D9",
    )
    add_docx_table(document, ["Tiempo", "Responsable", "Pantalla", "Objetivo"], [
        ["0:00-1:00", "Equipo", "Portada", "Presentacion, curso, grupo y objetivo."],
        ["1:00-3:00", "Angel", "Arquitectura", "Problema inicial y migracion gradual."],
        ["3:00-6:00", "Angel", "Repositorio/Gateway", "Estructura, rutas, JWT y ajuste frontend."],
        ["6:00-10:00", "Jussel", "Servicios y entidades", "Citas, Atencion, Caja, Notificaciones y Auth."],
        ["10:00-12:00", "Jussel", "Diagrama de BD", "Tres motores y cinco bases independientes."],
        ["12:00-15:00", "Juan", "Docker Compose", "Imagenes, contenedores, healthchecks y validacion."],
        ["15:00-19:00", "Juan", "Kubernetes", "Pods, Services, PVC, ConfigMap, Secret y port-forward."],
        ["19:00-22:00", "Equipo", "Postman", "Login y flujo Cita-Atencion-Deuda-Pago-Notificacion."],
        ["22:00-24:00", "Angel", "Frontend/Vercel", "Prueba de que /api entra por Gateway."],
        ["24:00-26:00", "Equipo", "Cierre", "Resultados, limitaciones y siguientes pasos."],
    ], [2.2, 3.2, 4, 7.3])

    document.add_heading("14.1. Texto sugerido: introduccion", level=2)
    add_paragraph(
        document,
        "Buenos dias. Somos el grupo 11 y presentamos la evolucion de Luz del Tunel, nuestro sistema de Clinica Privada, hacia una arquitectura basada en microservicios. Partimos de un backend Spring Boot centralizado y separamos tres dominios de negocio: Citas, Atencion Medica y Caja/Facturacion. Cada dominio posee su propia base de datos y utilizamos MongoDB, PostgreSQL y MySQL. Tambien incorporamos Gateway, Auth, Config Server, Eureka y Notificaciones. Mostraremos el codigo, el despliegue con Docker Compose y Kubernetes, y un flujo funcional completo."
    )

    document.add_heading("14.2. Texto sugerido: arquitectura", level=2)
    add_paragraph(
        document,
        "La aplicacion Angular no conoce las URLs internas. Todas las solicitudes entran al API Gateway. Para las APIs nuevas usamos el prefijo /api/ms; el Gateway lo elimina y reenvia la solicitud al servicio correspondiente. Las rutas que todavia pertenecen al backend original utilizan un fallback. Elegimos esta migracion gradual porque los contratos antiguos no son identicos a los nuevos y queriamos conservar el funcionamiento mientras extraemos capacidades."
    )

    document.add_heading("14.3. Texto sugerido: bases de datos", level=2)
    add_paragraph(
        document,
        "Cada microservicio es propietario de sus datos. Citas usa MongoDB por la naturaleza documental de la agenda. Atencion usa PostgreSQL por consistencia y estructura clinica. Caja usa MySQL por sus operaciones transaccionales y montos decimales. Auth tiene otra base MySQL y Notificaciones otra PostgreSQL. Compartir un motor no significa compartir una base; ningun servicio consulta directamente la persistencia de otro."
    )

    document.add_heading("14.4. Texto sugerido: Docker y Kubernetes", level=2)
    add_paragraph(
        document,
        "Docker empaqueta cada servicio y sus dependencias. Docker Compose levanta todo el entorno en un solo comando y respeta healthchecks. Kubernetes utiliza las mismas imagenes, pero declara Deployments, Services, volumenes, Secrets, ConfigMaps y readiness probes. En nuestro caso el cluster es local mediante Docker Desktop, lo cual cumple la demostracion de orquestacion sin mezclarlo con Render."
    )

    document.add_heading("14.5. Texto sugerido: punto 4", level=2)
    add_paragraph(
        document,
        "El frontend conserva /api como ruta base. En local, proxy.conf.json envia /api al Gateway en el puerto 8090. En Vercel, vercel.json reescribe /api al Gateway desplegado en Render. Esto desacopla Angular de la topologia interna y permite cambiar destinos sin modificar cada componente."
    )

    document.add_heading("14.6. Demostracion exacta", level=2)
    for text in [
        "Mostrar todos los Pods 1/1 o todos los contenedores saludables.",
        "Abrir /actuator/health del Gateway y /gateway/routes.",
        "Ejecutar login y guardar el JWT.",
        "Crear una cita y mostrar su ID MongoDB.",
        "Registrar y cerrar la atencion en PostgreSQL.",
        "Crear deuda y pago en MySQL; mostrar saldo y comprobante.",
        "Crear la notificacion de pago.",
        "Intentar duplicar el horario y mostrar que el servicio lo bloquea.",
        "Mostrar el frontend y una solicitud /api en Network apuntando al Gateway.",
    ]:
        add_numbered(document, text)

    add_section_break(document)
    document.add_heading("15. Preguntas que podria hacer un senior", level=1)
    add_docx_table(document, ["Pregunta", "Respuesta defendible"], [
        ["Por que no reescribieron todo?", "Una extraccion total incrementaba riesgo y rompia contratos existentes. Aplicamos migracion gradual con Gateway y fallback."],
        ["Esto es realmente microservicios?", "Si para los servicios extraidos: procesos separados, builds independientes, APIs propias y bases propias. El fallback representa deuda de migracion reconocida."],
        ["Por que tres motores?", "Cumple el requisito y demuestra persistencia poliglota: documentos para Citas, relacional estructurado para Atencion y transaccional para Caja."],
        ["Como manejan transacciones distribuidas?", "El prototipo usa coordinacion secuencial. Una evolucion productiva aplicaria eventos, idempotencia, outbox y Saga."],
        ["Eureka participa en Render?", "No. Se usa en local academico; en Render las URLs son variables directas y Eureka se deshabilita."],
        ["Por que Config Server?", "Centraliza configuracion comun en local. En nube gratuita se priorizo configuracion por entorno para reducir componentes activos."],
        ["Que pasa si cae Gateway?", "Es un punto unico de entrada. Kubernetes puede aumentar replicas y un balanceador distribuirlas; el despliegue actual usa una replica por alcance academico."],
        ["Como escalan una base con un servicio?", "Cada Deployment puede escalar de forma independiente, pero las bases locales son single-instance. En nube se delega persistencia al proveedor administrado."],
        ["Que observabilidad tienen?", "Actuator health, logs, Eureka y scripts de estado. Faltan metricas y trazas centralizadas; estan declaradas como mejora."],
        ["Que pruebas existen?", "Pruebas unitarias del router, compilacion completa, build Angular y pruebas integrales automatizadas por PowerShell/Postman."],
        ["Por que un proxy propio?", "Para mantener una implementacion pequena y visible en el curso. Spring Cloud Gateway seria una evolucion razonable para filtros, resiliencia y balanceo mas avanzados."],
        ["Como evitan acoplar bases?", "Cada servicio recibe IDs externos y nunca usa el repositorio o conexion de otra base."],
    ], [6, 10.7])

    add_section_break(document)
    document.add_heading("16. Guia para un junior que se incorpora", level=1)
    document.add_heading("16.1. Requisitos", level=2)
    for text in [
        "Java 21 y Maven Wrapper incluido.",
        "Node compatible con Angular 21 y npm.",
        "Docker Desktop con Kubernetes habilitado cuando corresponda.",
        "PowerShell y kubectl.",
        "Postman para importar la coleccion.",
    ]:
        add_bullet(document, text)

    document.add_heading("16.2. Primer recorrido", level=2)
    for text in [
        "Leer README.md y PROPUESTA_MICROSERVICIOS.md.",
        "Abrir pom.xml raiz y reconocer los ocho modulos.",
        "Elegir un servicio y seguir Controller -> DTO -> Service -> Repository -> Entity.",
        "Leer GatewayRouter y ProxyService para entender el punto de entrada.",
        "Revisar database/ para comparar entidades con DDL.",
        "Levantar Docker Compose y ejecutar verify-microservices.ps1.",
        "Importar Postman y repetir el flujo manualmente.",
        "Solo despues levantar Kubernetes y comparar Pods/Services con Compose.",
    ]:
        add_numbered(document, text)

    document.add_heading("16.3. Como agregar una operacion", level=2)
    for text in [
        "Definir el contrato de entrada y salida en DTOs.",
        "Agregar validaciones Jakarta en el DTO.",
        "Implementar la regla en el Service, no en el Controller.",
        "Agregar o ajustar consultas del Repository.",
        "Exponer el endpoint desde Controller.",
        "Protegerlo con SecurityConfig y JWT si corresponde.",
        "Agregar la solicitud a Postman y al script de verificacion.",
        "Actualizar schema/data si cambia persistencia.",
        "Compilar el modulo y luego el reactor completo.",
    ]:
        add_numbered(document, text)

    document.add_heading("16.4. Diagnostico basico", level=2)
    add_docx_table(document, ["Sintoma", "Que revisar"], [
        ["401 Unauthorized", "Bearer token, JWT_SECRET comun, expiracion y login de Auth Service."],
        ["502 Gateway", "URL destino, servicio despierto, logs del Gateway y cuerpo JSON valido."],
        ["Pod 0/1", "readinessProbe, logs, Secret, URL de base y estado del PVC."],
        ["Frontend llama 8080", "proxy.conf.json, npm start y cache del navegador."],
        ["Vercel llama al backend viejo", "vercel.json publicado y redeploy posterior al push."],
        ["Cita duplicada", "medico, fecha, rango horario y estado no CANCELADA."],
        ["Pago rechazado", "estado de deuda, saldo restante y monto positivo."],
    ], [5.6, 11.1])

    add_section_break(document)
    document.add_heading("17. Pendientes reales antes de grabar", level=1)
    add_callout(
        document,
        "Estado actual",
        "El codigo y la documentacion estan preparados localmente. La entrega aun necesita versionado, redeploy y evidencias recientes.",
        "FCE4D6",
    )
    for text in [
        "Revisar los cambios de Frontend y Microservices Container.",
        "Crear commits separados y hacer push a main.",
        "Verificar que Render redeploye Gateway y servicios afectados.",
        "Las rutas heredadas usan CITAS_LEGACY_ROUTE_ENABLED, ATENCION_LEGACY_ROUTE_ENABLED y CAJA_LEGACY_ROUTE_ENABLED en false para compatibilidad del frontend; las pruebas nuevas usan /api/ms.",
        "Esperar el redeploy de Vercel y revisar Network en el navegador.",
        "Ejecutar la coleccion Postman contra Render.",
        "Levantar Kubernetes limpio, ejecutar status-k8s.ps1 y verify-k8s.ps1.",
        "Tomar capturas claras de Pods, Eureka, Render, Aiven, Atlas y Postman.",
        "Ensayar el guion con cada integrante y comprobar audio/camara.",
        "Grabar sin mostrar .env, secretos, passwords o cadenas de conexion completas.",
    ]:
        add_numbered(document, text)

    document.add_heading("17.1. Criterio de listo para el punto 4", level=2)
    add_docx_table(document, ["Criterio", "Listo cuando"], [
        ["Codigo", "proxy.conf.json y vercel.json apuntan al Gateway."],
        ["Versionado", "Los cambios aparecen en GitHub."],
        ["Nube", "Vercel desplegado realiza /api contra Render Gateway."],
        ["Prueba", "Login funciona y una ruta /api/ms responde con JWT."],
        ["Evidencia", "Network o Postman muestra la URL del Gateway."],
    ], [5, 11.7])

    document.add_heading("17.2. Criterio de listo para el punto 5", level=2)
    add_docx_table(document, ["Criterio", "Listo cuando"], [
        ["Camara y audio", "Se ve y escucha a los integrantes que participan."],
        ["Implementacion", "Se muestran carpetas y codigo clave, no solo diapositivas."],
        ["Despliegue", "Docker Compose y Kubernetes se explican con evidencia real."],
        ["Funcionamiento", "Se ejecuta un flujo de negocio completo."],
        ["Arquitectura", "Se explica Gateway, servicios, motores y propiedad de datos."],
        ["Honestidad tecnica", "Se mencionan fallback, limites y pendientes sin sobreafirmar."],
    ], [5, 11.7])

    add_section_break(document)
    document.add_heading("18. Glosario rapido", level=1)
    add_docx_table(document, ["Termino", "Significado en este proyecto"], [
        ["Microservicio", "Aplicacion desplegable con responsabilidad, API y persistencia propias."],
        ["API Gateway", "Punto unico que recibe y reenvia solicitudes."],
        ["Service Discovery", "Registro para localizar servicios sin depender de IP fija."],
        ["Config Server", "Servidor que publica configuracion externa."],
        ["JWT", "Token firmado que transporta identidad y rol."],
        ["Docker image", "Paquete inmutable del servicio y su runtime."],
        ["Container", "Instancia en ejecucion de una imagen."],
        ["Deployment", "Objeto Kubernetes que mantiene Pods disponibles."],
        ["Service", "Nombre y red estable para acceder a Pods."],
        ["PVC", "Solicitud de almacenamiento persistente."],
        ["Readiness probe", "Comprobacion que indica si un Pod puede recibir trafico."],
        ["Fallback", "Ruta temporal hacia el backend original cuando un contrato aun no fue migrado."],
        ["Persistencia poliglota", "Uso de distintos motores segun el dominio."],
    ], [5, 11.7])

    document.add_heading("19. Cierre sugerido", level=1)
    add_paragraph(
        document,
        "Con esta implementacion demostramos una separacion real de dominios, persistencia independiente, seguridad JWT, entrada unificada y dos mecanismos de despliegue local. La arquitectura conserva temporalmente el backend original para permitir una migracion segura. El siguiente avance tecnico seria reemplazar el fallback por contratos completamente migrados, agregar comunicacion asincrona y fortalecer pruebas, observabilidad y automatizacion de despliegues."
    )
    add_callout(
        document,
        "Frase final para el equipo",
        "La mejor defensa no es memorizar nombres de herramientas; es poder seguir una solicitud desde Angular, pasar por Gateway, entrar al servicio, ejecutar la regla y comprobar en que base quedo el dato.",
        "E2F0D9",
    )

    document.save(GUIDE_PATH)
    print(GUIDE_PATH)


if __name__ == "__main__":
    create_guide()
