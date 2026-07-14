from __future__ import annotations

from pathlib import Path
from textwrap import wrap

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY, TA_LEFT
from reportlab.lib.pagesizes import A4, landscape
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import cm
from reportlab.platypus import (
    Image as PdfImage,
    PageBreak,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs"
EVIDENCE = DOCS / "evidence"
OUTPUT = DOCS / "entregable"
EVIDENCE.mkdir(parents=True, exist_ok=True)
OUTPUT.mkdir(parents=True, exist_ok=True)

ARCHITECTURE_IMAGE = EVIDENCE / "01-arquitectura-microservicios.png"
DATABASE_IMAGE = EVIDENCE / "02-esquemas-bases-datos.png"
VALIDATION_IMAGE = EVIDENCE / "03-validacion-local.png"
DOCX_PATH = OUTPUT / "261.Informe-Proyecto-Microservicios-DSW.docx"
PDF_PATH = OUTPUT / "261.Informe-Proyecto-Microservicios-DSW.pdf"

BLUE = "1F4E78"
TEAL = "087E8B"
GREEN = "2E7D32"
ORANGE = "C55A11"
GRAY = "5B6573"
LIGHT_BLUE = "D9EAF7"
LIGHT_GREEN = "E2F0D9"
LIGHT_ORANGE = "FCE4D6"
LIGHT_GRAY = "EEF1F4"


def font(size: int, bold: bool = False) -> ImageFont.FreeTypeFont:
    candidates = [
        Path("C:/Windows/Fonts/segoeuib.ttf" if bold else "C:/Windows/Fonts/segoeui.ttf"),
        Path("C:/Windows/Fonts/arialbd.ttf" if bold else "C:/Windows/Fonts/arial.ttf"),
    ]
    for candidate in candidates:
        if candidate.exists():
            return ImageFont.truetype(str(candidate), size)
    return ImageFont.load_default()


def multiline_center(draw: ImageDraw.ImageDraw, box: tuple[int, int, int, int], text: str,
                     fill: str, size: int, bold: bool = False, spacing: int = 8) -> None:
    selected = font(size, bold)
    width = box[2] - box[0] - 28
    words = text.split()
    lines: list[str] = []
    current = ""
    for word in words:
        candidate = f"{current} {word}".strip()
        if draw.textbbox((0, 0), candidate, font=selected)[2] <= width:
            current = candidate
        else:
            if current:
                lines.append(current)
            current = word
    if current:
        lines.append(current)
    line_height = selected.getbbox("Ag")[3] - selected.getbbox("Ag")[1]
    total = len(lines) * line_height + max(0, len(lines) - 1) * spacing
    y = box[1] + ((box[3] - box[1] - total) / 2)
    for line in lines:
        bbox = draw.textbbox((0, 0), line, font=selected)
        x = box[0] + ((box[2] - box[0] - (bbox[2] - bbox[0])) / 2)
        draw.text((x, y), line, font=selected, fill=fill)
        y += line_height + spacing


def rounded_box(draw: ImageDraw.ImageDraw, box: tuple[int, int, int, int], fill: str,
                outline: str, title: str, subtitle: str = "") -> None:
    draw.rounded_rectangle(box, radius=18, fill=fill, outline=outline, width=4)
    if subtitle:
        title_box = (box[0] + 12, box[1] + 14, box[2] - 12, box[1] + 70)
        subtitle_box = (box[0] + 12, box[1] + 65, box[2] - 12, box[3] - 10)
        multiline_center(draw, title_box, title, outline, 26, True, 4)
        multiline_center(draw, subtitle_box, subtitle, "#263238", 19, False, 4)
    else:
        multiline_center(draw, box, title, outline, 25, True, 4)


def arrow(draw: ImageDraw.ImageDraw, start: tuple[int, int], end: tuple[int, int], color: str = "#52606D") -> None:
    draw.line((start, end), fill=color, width=5)
    x, y = end
    if abs(end[1] - start[1]) >= abs(end[0] - start[0]):
        direction = 1 if end[1] > start[1] else -1
        points = [(x, y), (x - 11, y - 18 * direction), (x + 11, y - 18 * direction)]
    else:
        direction = 1 if end[0] > start[0] else -1
        points = [(x, y), (x - 18 * direction, y - 11), (x - 18 * direction, y + 11)]
    draw.polygon(points, fill=color)


def generate_architecture_image() -> None:
    image = Image.new("RGB", (2400, 1500), "#FFFFFF")
    draw = ImageDraw.Draw(image)
    draw.text((90, 45), "Arquitectura de Microservicios - Clinica Privada", font=font(48, True), fill="#17324D")
    draw.text((90, 108), "Entrada unica, servicios independientes y persistencia por dominio", font=font(27), fill="#52606D")

    frontend = (760, 180, 1640, 320)
    gateway = (760, 400, 1640, 555)
    rounded_box(draw, frontend, "#E8F1FB", "#1F4E78", "Frontend Angular", "Local :4200 / Vercel")
    rounded_box(draw, gateway, "#DCEFF2", "#087E8B", "API Gateway", "Entrada /api - aliases /api/ms - puerto 8090")
    arrow(draw, (1200, 320), (1200, 395))

    infra_left = (80, 400, 610, 555)
    infra_right = (1790, 400, 2320, 555)
    rounded_box(draw, infra_left, "#EEF1F4", "#5B6573", "Config Server", "Configuracion centralizada - puerto 8888")
    rounded_box(draw, infra_right, "#EEF1F4", "#5B6573", "Eureka Server", "Registro y descubrimiento - puerto 8761")
    arrow(draw, (610, 478), (750, 478))
    arrow(draw, (1790, 478), (1650, 478))

    services = [
        ((70, 690, 445, 890), "Auth Service", "JWT, usuarios y roles\nInfraestructura - :8091", "#E8F1FB", "#1F4E78"),
        ((530, 690, 905, 890), "Citas Service", "Reservas y disponibilidad\nNegocio - :8093", "#E2F0D9", "#2E7D32"),
        ((990, 690, 1410, 890), "Atencion Medica", "Diagnosticos y tratamientos\nNegocio - :8094", "#E2F0D9", "#2E7D32"),
        ((1495, 690, 1870, 890), "Caja Facturacion", "Deudas, pagos y comprobantes\nNegocio - :8095", "#E2F0D9", "#2E7D32"),
        ((1955, 690, 2330, 890), "Notificaciones", "Avisos multicanal\nApoyo - :8092", "#FCE4D6", "#C55A11"),
    ]
    for box, title, subtitle, fill, border in services:
        rounded_box(draw, box, fill, border, title, subtitle)
        arrow(draw, (1200, 555), ((box[0] + box[2]) // 2, box[1] - 8))

    databases = [
        ((90, 1040, 425, 1190), "MySQL", "clinica_auth_db", "#E8F1FB", "#1F4E78"),
        ((550, 1040, 885, 1190), "MongoDB", "clinica_citas_db", "#E2F0D9", "#2E7D32"),
        ((1032, 1040, 1368, 1190), "PostgreSQL", "clinica_atencion_db", "#E2F0D9", "#2E7D32"),
        ((1515, 1040, 1850, 1190), "MySQL", "clinica_caja_db", "#E2F0D9", "#2E7D32"),
        ((1975, 1040, 2310, 1190), "PostgreSQL", "clinica_notificaciones_db", "#FCE4D6", "#C55A11"),
    ]
    for index, (box, title, subtitle, fill, border) in enumerate(databases):
        rounded_box(draw, box, fill, border, title, subtitle)
        service_box = services[index][0]
        arrow(draw, ((service_box[0] + service_box[2]) // 2, service_box[3]), ((box[0] + box[2]) // 2, box[1] - 8))

    legacy = (760, 1285, 1640, 1415)
    rounded_box(draw, legacy, "#F7F7F7", "#7A7A7A", "Backend original", "Fallback temporal para modulos aun no migrados")
    draw.line((1640, 530, 2360, 620, 2360, 1350), fill="#8A8A8A", width=5)
    arrow(draw, (2360, 1350), (1650, 1350), "#8A8A8A")
    image.save(ARCHITECTURE_IMAGE, quality=95)


def draw_db_panel(draw: ImageDraw.ImageDraw, box: tuple[int, int, int, int], engine: str,
                  database: str, tables: list[tuple[str, list[str]]], color: str) -> None:
    draw.rounded_rectangle(box, radius=18, fill="#FFFFFF", outline=color, width=4)
    draw.rectangle((box[0], box[1], box[2], box[1] + 82), fill=color)
    draw.text((box[0] + 24, box[1] + 15), engine, font=font(28, True), fill="white")
    draw.text((box[0] + 24, box[1] + 49), database, font=font(20), fill="white")
    y = box[1] + 108
    for table_name, fields in tables:
        draw.text((box[0] + 25, y), table_name, font=font(24, True), fill=color)
        y += 38
        for field in fields:
            draw.text((box[0] + 40, y), field, font=font(18), fill="#263238")
            y += 28
        y += 16


def generate_database_image() -> None:
    image = Image.new("RGB", (2400, 1500), "#F7F9FB")
    draw = ImageDraw.Draw(image)
    draw.text((80, 45), "Esquemas de Bases de Datos por Microservicio", font=font(46, True), fill="#17324D")
    draw.text((80, 105), "Tres motores, cinco bases independientes y propiedad de datos por servicio", font=font(26), fill="#52606D")

    draw_db_panel(draw, (70, 180, 760, 760), "MySQL", "clinica_auth_db", [
        ("roles", ["PK id", "nombre (UNIQUE)", "descripcion", "activo"]),
        ("usuarios", ["PK id", "username / email (UNIQUE)", "password_hash", "FK rol_id", "created_at / updated_at"]),
    ], "#1F4E78")
    draw_db_panel(draw, (855, 180, 1545, 760), "MongoDB", "clinica_citas_db", [
        ("citas", ["_id (ObjectId)", "pacienteId / medicoId", "fecha / horaInicio / horaFin", "consultorio / motivo", "estado", "createdAt / updatedAt"]),
        ("indices", ["medicoId + fecha", "pacienteId", "estado"]),
    ], "#2E7D32")
    draw_db_panel(draw, (1640, 180, 2330, 760), "PostgreSQL", "clinica_atencion_db", [
        ("atenciones_medicas", ["PK id", "historia_clinica_id", "paciente_id / medico_id", "cita_id", "diagnostico / tratamiento", "estado", "created_at / updated_at"]),
    ], "#087E8B")
    draw_db_panel(draw, (450, 840, 1150, 1410), "MySQL", "clinica_caja_db", [
        ("deudas", ["PK id", "paciente_id", "concepto / referencia", "monto_total / monto_pagado", "estado"]),
        ("pagos", ["PK id", "deuda_id / paciente_id", "monto / metodo_pago", "numero_comprobante", "fecha_pago"]),
    ], "#C55A11")
    draw_db_panel(draw, (1250, 840, 1950, 1410), "PostgreSQL", "clinica_notificaciones_db", [
        ("notificaciones", ["PK id", "destinatario_tipo / id", "tipo / canal / estado", "titulo / mensaje", "referencia_tipo / id", "fechas de ciclo de vida"]),
    ], "#7B4F9D")
    image.save(DATABASE_IMAGE, quality=95)


def generate_validation_image() -> None:
    image = Image.new("RGB", (2200, 1300), "#FFFFFF")
    draw = ImageDraw.Draw(image)
    draw.text((80, 50), "Evidencia de validacion tecnica", font=font(46, True), fill="#17324D")
    draw.text((80, 110), "Resultados obtenidos durante la implementacion", font=font(26), fill="#52606D")

    panels = [
        ("Kubernetes local", [
            "Gateway health: UP",
            "Login por Gateway: OK",
            "Notificacion creada: OK",
            "Cita creada: OK",
            "Atencion registrada y cerrada: OK",
            "Deuda y pago registrados: OK",
            "Cruce de horario bloqueado: OK",
            "Validacion completa: OK",
        ]),
        ("Gateway Service", [
            "mvnw.cmd test",
            "Tests run: 3",
            "Failures: 0",
            "Errors: 0",
            "BUILD SUCCESS",
        ]),
        ("Frontend Angular", [
            "npm run build",
            "Environment API URL: /api",
            "Application bundle generation complete",
            "Salida: dist/frontend/browser",
        ]),
        ("Artefactos verificables", [
            "Docker Compose y Kubernetes",
            "Scripts PowerShell de arranque y prueba",
            "Coleccion Postman local y Render",
            "DDL/SQL/JS para tres motores",
        ]),
    ]
    coords = [(80, 200, 1050, 700), (1150, 200, 2120, 700), (80, 760, 1050, 1220), (1150, 760, 2120, 1220)]
    colors_panel = ["#087E8B", "#1F4E78", "#2E7D32", "#C55A11"]
    for (title, lines), box, color in zip(panels, coords, colors_panel):
        draw.rounded_rectangle(box, radius=16, fill="#F8FAFC", outline=color, width=4)
        draw.rectangle((box[0], box[1], box[2], box[1] + 72), fill=color)
        draw.text((box[0] + 24, box[1] + 17), title, font=font(27, True), fill="white")
        y = box[1] + 105
        for line in lines:
            draw.ellipse((box[0] + 28, y + 7, box[0] + 42, y + 21), fill=color)
            draw.text((box[0] + 60, y), line, font=font(22), fill="#263238")
            y += 45
    image.save(VALIDATION_IMAGE, quality=95)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    tc_pr.append(shading)


def set_cell_text(cell, text: str, bold: bool = False, color: str = "000000", size: float = 9.5) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "Arial"
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_docx_table(document: Document, headers: list[str], rows: list[list[str]], widths: list[float] | None = None):
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for index, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[index], header, True, "FFFFFF")
        set_cell_shading(table.rows[0].cells[index], BLUE)
    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            set_cell_text(cells[index], value)
            if len(table.rows) % 2 == 0:
                set_cell_shading(cells[index], "F5F8FA")
    if widths:
        for row in table.rows:
            for index, width in enumerate(widths):
                row.cells[index].width = Cm(width)
    document.add_paragraph()
    return table


def add_bullet(document: Document, text: str) -> None:
    paragraph = document.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.space_after = Pt(3)
    paragraph.add_run(text)


def add_numbered(document: Document, text: str) -> None:
    paragraph = document.add_paragraph(style="List Number")
    paragraph.paragraph_format.space_after = Pt(3)
    paragraph.add_run(text)


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    fld_char_1 = OxmlElement("w:fldChar")
    fld_char_1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char_2 = OxmlElement("w:fldChar")
    fld_char_2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_1)
    run._r.append(instr_text)
    run._r.append(fld_char_2)


def configure_docx(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.4)
    section.right_margin = Cm(2.4)
    add_page_number(section.footer.paragraphs[0])
    styles = document.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(10.5)
    styles["Normal"].paragraph_format.space_after = Pt(6)
    for name, size, color in [("Title", 20, BLUE), ("Heading 1", 15, BLUE), ("Heading 2", 12.5, TEAL), ("Heading 3", 11, GREEN)]:
        styles[name].font.name = "Arial"
        styles[name].font.size = Pt(size)
        styles[name].font.bold = True
        styles[name].font.color.rgb = RGBColor.from_string(color)


def create_docx() -> None:
    document = Document()
    configure_docx(document)

    for _ in range(2):
        document.add_paragraph()
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("UNIVERSIDAD NACIONAL MAYOR DE SAN MARCOS")
    r.bold = True
    r.font.name = "Arial"
    r.font.size = Pt(15)
    for text in [
        "Universidad del Peru, DECANA DE AMERICA",
        "FACULTAD DE INGENIERIA DE SISTEMAS E INFORMATICA",
        "Escuela Profesional de Ingenieria de Sistemas",
    ]:
        p = document.add_paragraph(text)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.add_paragraph()
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("APLICACION BASADA EN MICROSERVICIOS TOMANDO COMO BASE\nEL BACKEND DEL PROYECTO DEL CURSO")
    r.bold = True
    r.font.name = "Arial"
    r.font.size = Pt(18)
    r.font.color.rgb = RGBColor.from_string(BLUE)
    p = document.add_paragraph("Informe")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].bold = True
    p.runs[0].font.size = Pt(15)
    document.add_paragraph()
    add_docx_table(document, ["Dato", "Descripcion"], [
        ["PROFESORA", "Mamani Rodriguez, Zoraida Emperatriz"],
        ["CURSO", "Desarrollo de Sistemas Web"],
        ["GRUPO", "11"],
        ["INTEGRANTES", "Cotrina Ruiz, Angel Eduardo; Moya Laureano, Jussel Enrique; Tocto Caqui, Juan Andres"],
    ], [4, 12])
    for _ in range(3):
        document.add_paragraph()
    p = document.add_paragraph("Lima - Peru\n2026 - I")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.add_page_break()

    document.add_heading("Indice", level=1)
    for text in [
        "1. Equipo del Proyecto",
        "2. Arquitectura de Microservicios del Proyecto",
        "3. Entregables del Proyecto",
        "3.1. Microservicios",
        "3.1.1. URL del Repositorio Github",
        "3.1.2. Esquemas de las bases de datos",
        "3.1.3. Microservicios desplegados localmente",
        "3.1.4. URL de los microservicios desplegados en la nube",
        "3.1.5. Coleccion de servicios definidos en Postman",
        "3.2. Frontend",
        "3.2.1. URL del Repositorio Github",
        "4. Logros del Sprint",
        "5. Pendientes del Sprint",
        "6. Conclusiones",
        "7. Recomendaciones",
        "8. Anexos",
    ]:
        document.add_paragraph(text)
    document.add_page_break()

    document.add_heading("1. Equipo del Proyecto", level=1)
    add_docx_table(document, ["Nombre", "Rol", "Cuenta Github"], [
        ["Cotrina Ruiz, Angel Eduardo", "Desarrollador principal / Full Stack", "AngelCotrina28"],
        ["Moya Laureano, Jussel Enrique", "Scrum Master / Full Stack", "Enrike352"],
        ["Tocto Caqui, Juan Andres", "Product Owner / Full Stack", "RexoJuan"],
    ], [5, 7, 4])

    document.add_heading("2. Arquitectura de Microservicios del Proyecto", level=1)
    document.add_paragraph(
        "La propuesta aplica una migracion gradual desde el backend original hacia servicios autonomos. "
        "El frontend utiliza el API Gateway como punto unico de entrada. Los endpoints nuevos se exponen "
        "mediante el prefijo /api/ms, mientras que el Gateway conserva un fallback temporal para modulos "
        "que aun pertenecen al backend original."
    )
    document.add_picture(str(ARCHITECTURE_IMAGE), width=Inches(6.6))
    document.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_docx_table(document, ["Componente", "Tipo", "Responsabilidad"], [
        ["API Gateway", "Infraestructura", "Entrada unica, CORS, propagacion de JWT y enrutamiento."],
        ["Config Server", "Infraestructura", "Centraliza configuracion para ejecucion local."],
        ["Eureka Server", "Infraestructura", "Registra y descubre instancias en el entorno local."],
        ["Auth Service", "Infraestructura", "Autentica usuarios y emite tokens JWT."],
        ["Citas Service", "Negocio", "Gestiona reserva, disponibilidad y cancelacion de citas."],
        ["Atencion Medica Service", "Negocio", "Registra diagnosticos, tratamientos y cierre de atenciones."],
        ["Caja Facturacion Service", "Negocio", "Gestiona deudas, pagos, saldos y comprobantes."],
        ["Notificaciones Service", "Apoyo", "Registra y consulta avisos para pacientes y trabajadores."],
    ], [4.5, 3.2, 9])

    document.add_heading("3. Entregables del Proyecto", level=1)
    document.add_heading("3.1. Microservicios", level=2)
    document.add_paragraph(
        "El repositorio contenedor incluye ocho aplicaciones Spring Boot, cada una con su propio pom.xml "
        "y Dockerfile. Los tres servicios de negocio emplean motores distintos: MongoDB, PostgreSQL y MySQL."
    )

    document.add_heading("3.1.1. URL del Repositorio Github", level=3)
    document.add_paragraph("https://github.com/AngelCotrina28/ClinicaPrivadaMSContainer")
    document.add_paragraph(
        "El historial Git evidencia incorporacion incremental de Gateway, bases en nube, Config Server, "
        "Eureka, Docker Compose y Kubernetes."
    )

    document.add_heading("3.1.2. Esquemas de las bases de datos", level=3)
    document.add_picture(str(DATABASE_IMAGE), width=Inches(6.6))
    document.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_docx_table(document, ["Servicio", "Motor", "Base", "Artefactos"], [
        ["Auth", "MySQL", "clinica_auth_db", "database/mysql/auth/schema.sql y data.sql"],
        ["Citas", "MongoDB", "clinica_citas_db", "database/mongodb/citas/*.mongodb.js"],
        ["Atencion Medica", "PostgreSQL", "clinica_atencion_db", "database/postgresql/atencion/*.sql"],
        ["Caja Facturacion", "MySQL", "clinica_caja_db", "database/mysql/caja/*.sql"],
        ["Notificaciones", "PostgreSQL", "clinica_notificaciones_db", "database/postgresql/notificaciones/*.sql"],
    ], [3.5, 2.8, 4.2, 6])

    document.add_heading("3.1.3. Microservicios desplegados localmente", level=3)
    document.add_paragraph("Se implementaron dos mecanismos locales equivalentes:")
    add_bullet(document, "Docker Compose: docker-compose.microservices.yml y scripts/start-microservices.ps1.")
    add_bullet(document, "Kubernetes: cinco manifiestos ordenados en k8s/ y scripts/start-k8s.ps1.")
    add_bullet(document, "Verificacion: scripts/verify-microservices.ps1 y scripts/verify-k8s.ps1.")
    document.add_paragraph("Orden de despliegue: bases de datos -> Config Server -> Eureka -> microservicios -> Gateway.")
    document.add_picture(str(VALIDATION_IMAGE), width=Inches(6.6))
    document.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    document.add_heading("3.1.4. URL de los Microservicios Desplegados en la nube", level=3)
    add_docx_table(document, ["Servicio", "URL publica", "Persistencia"], [
        ["Gateway", "https://clinica-gateway-service.onrender.com", "No aplica"],
        ["Auth", "https://clinica-auth-service.onrender.com", "MySQL en Aiven"],
        ["Citas", "https://clinica-citas-service.onrender.com", "MongoDB Atlas"],
        ["Atencion Medica", "https://clinica-atencion-medica-service.onrender.com", "PostgreSQL en Aiven"],
        ["Caja Facturacion", "https://clinica-caja-facturacion-service.onrender.com", "MySQL en Aiven"],
        ["Notificaciones", "https://clinica-notificaciones-service.onrender.com", "PostgreSQL en Aiven"],
    ], [3.5, 8.3, 4.2])
    document.add_paragraph(
        "Los servicios del plan gratuito pueden entrar en reposo. Antes de la demostracion se recomienda "
        "abrir /actuator/health y esperar la respuesta {status: UP}."
    )

    document.add_heading("3.1.5. Coleccion de servicios definidos en Postman", level=3)
    document.add_paragraph(
        "La carpeta postman/ contiene una coleccion Postman 2.1 y dos ambientes. El flujo guarda "
        "automaticamente el JWT y los identificadores de cita, atencion, deuda, pago y notificacion."
    )
    add_bullet(document, "postman/Clinica-Privada-Microservicios.postman_collection.json")
    add_bullet(document, "postman/Clinica-Local.postman_environment.json")
    add_bullet(document, "postman/Clinica-Render.postman_environment.json")
    document.add_paragraph(
        "La coleccion comprueba salud del Gateway, login, perfil, cita, disponibilidad, atencion, "
        "cierre medico, deuda, pago, resumen de caja, notificacion y cancelacion."
    )

    document.add_heading("3.2. Frontend", level=2)
    document.add_heading("3.2.1. URL del Repositorio Github", level=3)
    document.add_paragraph("https://github.com/AngelCotrina28/Clinica-Privada-Frontend")
    document.add_paragraph("Frontend desplegado: https://clinica-privada-frontend.vercel.app/")
    document.add_paragraph(
        "El frontend Angular mantiene environment.apiUrl=/api. En desarrollo, proxy.conf.json dirige "
        "las solicitudes al Gateway local en http://localhost:8090. En Vercel, vercel.json dirige /api "
        "al Gateway de Render. De esta forma no se codifican URLs de microservicios en los componentes."
    )

    document.add_heading("4. Logros del Sprint", level=1)
    for text in [
        "Implementacion de tres microservicios de negocio con persistencia independiente y tres motores distintos.",
        "Implementacion de Auth, Notificaciones, Gateway, Config Server y Eureka Server.",
        "Despliegue local reproducible mediante Docker Compose y Kubernetes.",
        "Despliegue en Render conectado con Aiven y MongoDB Atlas.",
        "Propagacion de JWT Bearer desde Gateway hacia servicios protegidos.",
        "Ajuste del frontend para utilizar el Gateway en local y nube.",
        "Esquemas de datos versionados y coleccion Postman automatizada.",
        "Pruebas del Gateway: 3 ejecutadas, 0 fallos; compilacion Angular satisfactoria.",
    ]:
        add_bullet(document, text)

    document.add_heading("5. Pendientes del Sprint", level=1)
    document.add_paragraph(
        "Los requisitos principales quedaron implementados. Como trabajo evolutivo, no bloqueante para "
        "la entrega actual, se identifican los siguientes puntos:"
    )
    add_bullet(document, "Migrar gradualmente los endpoints heredados restantes y retirar el fallback al backend original.")
    add_bullet(document, "Incorporar CI/CD para ejecutar pruebas y publicar imagenes de forma automatica.")
    add_bullet(document, "Desplegar Kubernetes en un proveedor administrado si se requiere un entorno K8s publico.")
    add_bullet(document, "Agregar observabilidad centralizada de metricas, trazas y logs.")

    document.add_heading("6. Conclusiones", level=1)
    add_numbered(document, "La separacion por dominios reduce el acoplamiento y permite desplegar cada capacidad de negocio de manera independiente.")
    add_numbered(document, "El uso de MongoDB, PostgreSQL y MySQL evidencia persistencia poliglota basada en la naturaleza de cada dominio.")
    add_numbered(document, "Docker Compose simplifica el arranque integral y Kubernetes agrega declaracion de estado, servicios internos y recuperacion de contenedores.")
    add_numbered(document, "El API Gateway mantiene una interfaz unica para Angular y permite una migracion gradual sin modificar todas las pantallas simultaneamente.")

    document.add_heading("7. Recomendaciones", level=1)
    add_numbered(document, "Ejecutar verify-k8s.ps1 o la coleccion Postman antes de cada demostracion.")
    add_numbered(document, "No versionar contrasenas; configurarlas como variables de entorno o Secrets de Kubernetes.")
    add_numbered(document, "Mantener la propiedad exclusiva de cada base de datos y evitar consultas cruzadas entre servicios.")
    add_numbered(document, "Crear contratos versionados para sustituir progresivamente los endpoints del backend original.")

    document.add_heading("8. Anexos", level=1)
    document.add_heading("8.1. Contribuciones obtenidas de Git", level=2)
    add_docx_table(document, ["Repositorio", "Cuenta", "Commits registrados"], [
        ["Microservices Container", "AngelCotrina28", "16"],
        ["Frontend", "AngelCotrina28 / Angel Cotrina", "42"],
        ["Frontend", "RexoJuan (autor Git: JuanTD / Juan TD)", "30"],
        ["Frontend", "Enrike352", "26"],
        ["Backend", "AngelCotrina28 / Angel Cotrina", "65"],
        ["Backend", "RexoJuan (autor Git: JuanTD / Juan TD)", "20"],
        ["Backend", "Enrike352", "24"],
    ], [5, 6, 4])
    document.add_paragraph(
        "Los alias de autor fueron agrupados por correo y cuenta. La asignacion cualitativa final debe "
        "complementarse con la explicacion del equipo en la exposicion."
    )
    document.add_heading("8.2. Comandos de verificacion", level=2)
    for command in [
        ".\\scripts\\start-microservices.ps1",
        ".\\scripts\\verify-microservices.ps1",
        ".\\scripts\\start-k8s.ps1",
        ".\\scripts\\status-k8s.ps1",
        ".\\scripts\\verify-k8s.ps1",
        "npm.cmd run build",
        ".\\mvnw.cmd test",
    ]:
        p = document.add_paragraph()
        run = p.add_run(command)
        run.font.name = "Consolas"
        run.font.size = Pt(9.5)
    document.add_heading("8.3. Evidencias finales pendientes", level=2)
    add_bullet(document, "Capturas finales de Docker Desktop, Kubernetes, Render, Aiven y MongoDB Atlas para el video y la version definitiva.")

    document.save(DOCX_PATH)


def pdf_header_footer(canvas, doc) -> None:
    canvas.saveState()
    canvas.setStrokeColor(colors.HexColor("#D0D7DE"))
    canvas.line(2 * cm, 1.5 * cm, A4[0] - 2 * cm, 1.5 * cm)
    canvas.setFont("Helvetica", 8)
    canvas.setFillColor(colors.HexColor("#52606D"))
    canvas.drawCentredString(A4[0] / 2, 1.05 * cm, f"Proyecto de Microservicios - Pagina {doc.page}")
    canvas.restoreState()


def create_pdf() -> None:
    styles = getSampleStyleSheet()
    body = ParagraphStyle("Body", parent=styles["BodyText"], fontName="Helvetica", fontSize=9.4,
                          leading=13, alignment=TA_JUSTIFY, spaceAfter=7)
    h1 = ParagraphStyle("H1", parent=styles["Heading1"], fontName="Helvetica-Bold", fontSize=15,
                        textColor=colors.HexColor(f"#{BLUE}"), spaceBefore=10, spaceAfter=8)
    h2 = ParagraphStyle("H2", parent=styles["Heading2"], fontName="Helvetica-Bold", fontSize=12,
                        textColor=colors.HexColor(f"#{TEAL}"), spaceBefore=8, spaceAfter=6)
    h3 = ParagraphStyle("H3", parent=styles["Heading3"], fontName="Helvetica-Bold", fontSize=10.5,
                        textColor=colors.HexColor(f"#{GREEN}"), spaceBefore=6, spaceAfter=5)
    centered = ParagraphStyle("Centered", parent=body, alignment=TA_CENTER)
    cover = ParagraphStyle("Cover", parent=centered, fontName="Helvetica-Bold", fontSize=17,
                           leading=22, textColor=colors.HexColor(f"#{BLUE}"), spaceAfter=18)
    small = ParagraphStyle("Small", parent=body, fontSize=8, leading=10)

    doc = SimpleDocTemplate(str(PDF_PATH), pagesize=A4, rightMargin=2 * cm, leftMargin=2 * cm,
                            topMargin=1.8 * cm, bottomMargin=2 * cm, title="Informe Proyecto Microservicios DSW")
    story = []
    story.extend([
        Spacer(1, 2 * cm),
        Paragraph("UNIVERSIDAD NACIONAL MAYOR DE SAN MARCOS", ParagraphStyle("Uni", parent=cover, fontSize=15)),
        Paragraph("Universidad del Peru, DECANA DE AMERICA", centered),
        Paragraph("FACULTAD DE INGENIERIA DE SISTEMAS E INFORMATICA", centered),
        Paragraph("Escuela Profesional de Ingenieria de Sistemas", centered),
        Spacer(1, 1.3 * cm),
        Paragraph("APLICACION BASADA EN MICROSERVICIOS TOMANDO COMO BASE<br/>EL BACKEND DEL PROYECTO DEL CURSO", cover),
        Paragraph("Informe", ParagraphStyle("Report", parent=cover, fontSize=14)),
        Spacer(1, 0.8 * cm),
    ])
    cover_table = Table([
        ["PROFESORA", "Mamani Rodriguez, Zoraida Emperatriz"],
        ["CURSO", "Desarrollo de Sistemas Web"],
        ["GRUPO", "11"],
        ["INTEGRANTES", "Cotrina Ruiz, Angel Eduardo; Moya Laureano, Jussel Enrique; Tocto Caqui, Juan Andres"],
    ], colWidths=[4 * cm, 11 * cm])
    cover_table.setStyle(TableStyle([
        ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#B8C2CC")),
        ("BACKGROUND", (0, 0), (0, -1), colors.HexColor(f"#{LIGHT_BLUE}")),
        ("FONTNAME", (0, 0), (0, -1), "Helvetica-Bold"),
        ("FONTNAME", (0, 0), (-1, -1), "Helvetica"),
        ("FONTSIZE", (0, 0), (-1, -1), 9),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 7),
        ("TOPPADDING", (0, 0), (-1, -1), 7),
    ]))
    story.extend([cover_table, Spacer(1, 2.5 * cm), Paragraph("Lima - Peru<br/>2026 - I", centered), PageBreak()])

    def heading(text: str, level: int = 1):
        story.append(Paragraph(text, {1: h1, 2: h2, 3: h3}[level]))

    def para(text: str):
        story.append(Paragraph(text, body))

    def bullets(items: list[str]):
        for item in items:
            story.append(Paragraph(f"&#8226; {item}", body))

    def pdf_table(headers: list[str], rows: list[list[str]], widths: list[float]):
        data = [[Paragraph(h, small) for h in headers]] + [[Paragraph(str(v), small) for v in row] for row in rows]
        table = Table(data, colWidths=[w * cm for w in widths], repeatRows=1)
        table.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor(f"#{BLUE}")),
            ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
            ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
            ("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#B8C2CC")),
            ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#F5F8FA")]),
            ("LEFTPADDING", (0, 0), (-1, -1), 5),
            ("RIGHTPADDING", (0, 0), (-1, -1), 5),
            ("TOPPADDING", (0, 0), (-1, -1), 5),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
        ]))
        story.extend([table, Spacer(1, 0.25 * cm)])

    heading("Indice")
    para("1. Equipo del Proyecto<br/>2. Arquitectura de Microservicios del Proyecto<br/>3. Entregables del Proyecto<br/>"
         "3.1. Microservicios<br/>3.2. Frontend<br/>4. Logros del Sprint<br/>5. Pendientes del Sprint<br/>"
         "6. Conclusiones<br/>7. Recomendaciones<br/>8. Anexos")
    story.append(PageBreak())
    heading("1. Equipo del Proyecto")
    pdf_table(["Nombre", "Rol", "Github"], [
        ["Cotrina Ruiz, Angel Eduardo", "Desarrollador principal / Full Stack", "AngelCotrina28"],
        ["Moya Laureano, Jussel Enrique", "Scrum Master / Full Stack", "Enrike352"],
        ["Tocto Caqui, Juan Andres", "Product Owner / Full Stack", "RexoJuan"],
    ], [4.2, 7.5, 3.5])
    heading("2. Arquitectura de Microservicios del Proyecto")
    para("La propuesta aplica una migracion gradual. Angular utiliza el API Gateway como entrada unica; /api/ms selecciona los servicios nuevos y el fallback mantiene temporalmente los modulos heredados.")
    story.append(PdfImage(str(ARCHITECTURE_IMAGE), width=17 * cm, height=10.63 * cm))
    pdf_table(["Componente", "Tipo", "Responsabilidad"], [
        ["Gateway", "Infraestructura", "Entrada unica, JWT y enrutamiento."],
        ["Config Server", "Infraestructura", "Configuracion centralizada local."],
        ["Eureka", "Infraestructura", "Registro y descubrimiento local."],
        ["Auth", "Infraestructura", "Login, usuarios, roles y JWT."],
        ["Citas", "Negocio", "Reservas, disponibilidad y cancelacion."],
        ["Atencion Medica", "Negocio", "Diagnosticos y tratamientos."],
        ["Caja Facturacion", "Negocio", "Deudas, pagos y comprobantes."],
        ["Notificaciones", "Apoyo", "Avisos para pacientes y trabajadores."],
    ], [4.2, 3.3, 7.7])
    heading("3. Entregables del Proyecto")
    heading("3.1. Microservicios", 2)
    para("El repositorio contiene ocho aplicaciones Spring Boot independientes. Los tres servicios de negocio usan MongoDB, PostgreSQL y MySQL.")
    heading("3.1.1. URL del Repositorio Github", 3)
    para("https://github.com/AngelCotrina28/ClinicaPrivadaMSContainer")
    heading("3.1.2. Esquemas de las bases de datos", 3)
    story.append(PdfImage(str(DATABASE_IMAGE), width=17 * cm, height=10.63 * cm))
    pdf_table(["Servicio", "Motor", "Base", "Artefactos"], [
        ["Auth", "MySQL", "clinica_auth_db", "database/mysql/auth"],
        ["Citas", "MongoDB", "clinica_citas_db", "database/mongodb/citas"],
        ["Atencion", "PostgreSQL", "clinica_atencion_db", "database/postgresql/atencion"],
        ["Caja", "MySQL", "clinica_caja_db", "database/mysql/caja"],
        ["Notificaciones", "PostgreSQL", "clinica_notificaciones_db", "database/postgresql/notificaciones"],
    ], [3.2, 2.7, 4, 5.3])
    heading("3.1.3. Microservicios desplegados localmente", 3)
    bullets(["Docker Compose mediante start-microservices.ps1.", "Kubernetes mediante start-k8s.ps1.", "Validacion integral mediante verify-microservices.ps1 y verify-k8s.ps1."])
    story.append(PdfImage(str(VALIDATION_IMAGE), width=17 * cm, height=10.05 * cm))
    heading("3.1.4. URL de los Microservicios Desplegados en la nube", 3)
    pdf_table(["Servicio", "URL", "Base"], [
        ["Gateway", "clinica-gateway-service.onrender.com", "No aplica"],
        ["Auth", "clinica-auth-service.onrender.com", "MySQL / Aiven"],
        ["Citas", "clinica-citas-service.onrender.com", "MongoDB Atlas"],
        ["Atencion", "clinica-atencion-medica-service.onrender.com", "PostgreSQL / Aiven"],
        ["Caja", "clinica-caja-facturacion-service.onrender.com", "MySQL / Aiven"],
        ["Notificaciones", "clinica-notificaciones-service.onrender.com", "PostgreSQL / Aiven"],
    ], [3, 8.5, 3.7])
    para("En el plan gratuito los servicios pueden entrar en reposo; antes de exponer se debe abrir /actuator/health y esperar status UP.")
    heading("3.1.5. Coleccion de servicios definidos en Postman", 3)
    para("postman/ contiene la coleccion Postman 2.1 y ambientes local/Render. El flujo almacena automaticamente JWT e identificadores y prueba Citas, Atencion, Caja y Notificaciones.")
    heading("3.2. Frontend", 2)
    heading("3.2.1. URL del Repositorio Github", 3)
    para("https://github.com/AngelCotrina28/Clinica-Privada-Frontend")
    para("Frontend desplegado: https://clinica-privada-frontend.vercel.app/")
    para("Angular conserva /api como URL unica. proxy.conf.json apunta al Gateway local :8090 y vercel.json al Gateway de Render.")
    heading("4. Logros del Sprint")
    bullets([
        "Tres microservicios de negocio con tres motores distintos.",
        "Servicios de infraestructura y apoyo implementados.",
        "Docker Compose y Kubernetes locales reproducibles.",
        "Servicios desplegados en Render con bases de Aiven y MongoDB Atlas.",
        "JWT propagado por Gateway y frontend ajustado.",
        "DDL, datos de ejemplo y coleccion Postman versionados.",
        "Gateway: 3 pruebas, 0 fallos. Angular: build satisfactorio.",
    ])
    heading("5. Pendientes del Sprint")
    para("Como trabajo evolutivo no bloqueante quedan la migracion total del backend original, CI/CD, observabilidad y un despliegue Kubernetes administrado en nube.")
    heading("6. Conclusiones")
    bullets([
        "La separacion por dominios reduce acoplamiento y permite despliegues independientes.",
        "La persistencia poliglota asigna cada motor segun las necesidades del dominio.",
        "Docker Compose facilita el arranque y Kubernetes administra el estado declarado.",
        "El Gateway permite una migracion gradual sin reescribir Angular de una sola vez.",
    ])
    heading("7. Recomendaciones")
    bullets([
        "Ejecutar verify-k8s.ps1 o Postman antes de cada demostracion.",
        "Mantener contrasenas en variables o Secrets, nunca en Git.",
        "Evitar acceso directo entre bases de datos de distintos servicios.",
        "Versionar los contratos mientras se retira el fallback heredado.",
    ])
    heading("8. Anexos")
    heading("8.1. Contribuciones obtenidas de Git", 2)
    pdf_table(["Repositorio", "Cuenta", "Commits"], [
        ["Microservices Container", "AngelCotrina28", "16"],
        ["Frontend", "AngelCotrina28 / Angel Cotrina", "42"],
        ["Frontend", "RexoJuan (Git: JuanTD / Juan TD)", "30"],
        ["Frontend", "Enrike352", "26"],
        ["Backend", "AngelCotrina28 / Angel Cotrina", "65"],
        ["Backend", "RexoJuan (Git: JuanTD / Juan TD)", "20"],
        ["Backend", "Enrike352", "24"],
    ], [5.5, 7, 2.7])
    heading("8.2. Evidencias finales pendientes", 2)
    bullets(["Capturas finales de Docker Desktop, Kubernetes, Render, Aiven y MongoDB Atlas para la version entregable y el video."])
    doc.build(story, onFirstPage=pdf_header_footer, onLaterPages=pdf_header_footer)


def main() -> None:
    generate_architecture_image()
    generate_database_image()
    generate_validation_image()
    create_docx()
    create_pdf()
    print(ARCHITECTURE_IMAGE)
    print(DATABASE_IMAGE)
    print(VALIDATION_IMAGE)
    print(DOCX_PATH)
    print(PDF_PATH)


if __name__ == "__main__":
    main()
